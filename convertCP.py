"""
    author: Yihui "Ray" Ren
    email:  yren@bnl.gov
    description: convert NSF/DOE C&P pdf to plaintext table and NIH docx
    disclaimer: this code is neither supported nor approved by DOE, NSF or NIH

    C&P versions:
    DOE: SCV C&P(O)S v.2023-1 (rev.01/31/2023)
    NIH: OMB No. 0925-0001 and 0925-0002 (Rev. 12/2020 Approved Through 02/28/2023)
"""

from dataclasses import dataclass, field, asdict
from typing import List
from collections import OrderedDict
from pathlib import Path
import string
import argparse
import itertools as it

from pypdf import PdfReader
from tabulate import tabulate
from docx import Document
from docx.shared import Pt, Inches


@dataclass
class Project:
    """keep  project info organized"""

    title: str
    status_of_support: str
    award_number: str
    source_of_support: str
    place: str
    start_date: str
    end_date: str
    total_amount: int
    person_year_month: str = field(repr=False)
    objective: str
    overlap: str
    pym_year: List[int] = field(init=False, repr=True)
    pym_person: List[float] = field(init=False, repr=True)

    def __post_init__(self):
        tokens = self.person_year_month.split(" ")
        tokens = [x for x in tokens if x.startswith(tuple(string.digits))]
        years = tokens[::2]
        pmonth = tokens[1::2]
        self.pym_year = [int(x) for x in years]
        self.pym_person = [float(x) for x in pmonth]


def print_page(page):
    lines = page.extract_text().split("\n")
    print(repr(lines))


def merge_pages(reader):
    ans = []
    num_of_pages = len(reader.pages)
    for i in range(num_of_pages):
        page = reader.pages[i]
        lines = page.extract_text().split("\n")
        lines = [line.strip() for line in lines]
        ans += lines
    return ans


K = [  # `K` may change for every C&P format revision
    "*Project/Proposal Title:",
    "*Status of Support:",
    "Proposal/Award Number:",
    "*Source of Support:",
    "*Primary Place of Performance:",
    "*Project/Proposal Support Start Date: (MM/YYYY):",
    "*Project/Proposal Support End Date: (MM/YYYY):",
    "*Total Award Amount:",
    "* Person Months (Calendar/Academic/Summer) per budget period Committed to the",
    "*Overall Objectives:",
    "*Statement of Potential Overlap:",
]
ksz = len(K)


def filter_line(line):
    return line.startswith("SCV C&P") or line.startswith("Page")


def stop_line(line):
    return line.startswith("Certification:")


def consume_until(idx, sz, lines, stop_signal):
    rst = ""
    idx = idx + 1
    while idx < sz and not stop_signal(lines[idx]):
        if stop_line(lines[idx]):
            idx = sz
            continue
        if not filter_line(lines[idx]):
            rst += lines[idx] + " "
        idx += 1
    return rst.strip(), idx


def process_lines(lines):
    rst = list()
    sz = len(lines)
    i, j, buff = 0, 0, []
    while i < sz:
        line = lines[i]
        if line == K[j]:
            j += 1
            tmp, i = consume_until(i, sz, lines, lambda x: x == K[j % ksz])
            buff.append(tmp)
            if j == ksz:
                rst.append(Project(*buff))
                j, buff = 0, []
        else:
            i += 1

    if len(buff) > 0:
        rst.append(Project(*buff))

    return rst


def to_nih_format(proj):
    nih = OrderedDict()
    nih["*Title"] = proj.title
    nih["*Major Goals"] = proj.objective
    nih["*Status of Support"] = proj.status_of_support
    nih["Project Number"] = proj.award_number
    nih["Name of PD/PI"] = "TODO"  # NSF/DOE C&P does not have this field
    nih["*Source of Support"] = proj.source_of_support
    nih["*Primary Place of Performance"] = proj.place
    nih["Project/Proposal Start and End Date: (MM/YYYY) (if available)"] = (
        proj.start_date + " -- " + proj.end_date
    )
    nih["*Total Award Amount (including Indirect Costs)"] = proj.total_amount
    return nih


def add_table(doc, proj):
    doc.add_paragraph("*Person Months (Calendar/Academic/Summer) per budget period.")
    style = doc.styles["Normal"]
    table = doc.add_table(rows=1, cols=2)
    table.style = "TableGrid"
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Inches(1.62)
    table.columns[1].width = Inches(1.81)
    header = table.rows[0].cells
    header[0].text = "Year(YYYY)"
    header[1].text = "Person Months (##.##)"
    header[0].width = Inches(1.62)
    header[1].width = Inches(1.81)
    for yr, prsn, idx in zip(proj.pym_year, proj.pym_person, it.count(1)):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx) + ". " + str(yr)
        row_cells[1].text = f"{prsn:.2f}" + " calendar"
        idx += 1


def add_styled_paragraph(doc):
    doc.add_paragraph()
    style = doc.styles["Normal"]
    style.font.name, style.font.size = "Arial", Pt(11)
    style.paragraph_format.line_spacing = 1
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(6)


def output_nih_docx(projects):
    act_proj = [p for p in projects if p.status_of_support == "current"]
    pnd_proj = [p for p in projects if p.status_of_support == "pending"]
    assert len(act_proj) + len(pnd_proj) == len(
        projects
    ), "invalid project status of support"

    doc = Document()
    font = doc.styles["Normal"].font
    font.name, font.size = "Arial", Pt(11)

    doc.add_paragraph("ACTIVE")
    for proj in act_proj:
        add_styled_paragraph(doc)
        nih = to_nih_format(proj)
        for k, v in nih.items():
            doc.add_paragraph(k + ": " + v)

        add_table(doc, proj)
        doc.add_paragraph("  ")

    doc.add_paragraph("PENDING")
    for proj in pnd_proj:
        add_styled_paragraph(doc)
        nih = to_nih_format(proj)
        for k, v in nih.items():
            doc.add_paragraph(k + ": " + v)

        add_table(doc, proj)
        doc.add_paragraph("  ")

    return doc


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        prog="extract current and pending",
        description="convert NSF/DOE Current and Pending (rev.01/31/2023) to plain text table",
    )
    parser.add_argument("filename")
    parser.add_argument(
        "--nih",
        action="store_true",
        help="whether or not convert to NIH MS-word format",
    )
    args = parser.parse_args()

    reader = PdfReader(args.filename)
    lines = merge_pages(reader)
    projects = process_lines(lines)

    for proj in projects:
        print(tabulate(asdict(proj).items(), headers="keys"))

    if args.nih:
        doc = output_nih_docx(projects)
        doc.save("NIH_" + Path(args.filename).stem + ".docx")
